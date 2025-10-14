# app.py
import os
import re
from collections import Counter
from datetime import datetime
from flask import Flask, request, render_template_string, send_from_directory, redirect, url_for
from werkzeug.utils import secure_filename

import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize


nltk.download('punkt')
nltk.download('stopwords')

# Optional libs (may need installation)
try:
    import pdfplumber
except Exception:
    pdfplumber = None
try:
    import docx
except Exception:
    docx = None

import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

try:
    import spacy
except Exception:
    spacy = None




ALLOWED_EXT = {'pdf', 'txt', 'docx'}
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, 'uploads')
STATIC_DIR = os.path.join(BASE_DIR, 'static')
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_DIR
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024

INDEX_HTML = """
<!doctype html>
<html>
<head><meta charset="utf-8"><title>Document Analyzer</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light">
<div class="container py-4">
  <h1>Document Analyzer</h1>
  <p>Upload a PDF / DOCX / TXT or paste text below.</p>
  <form method="post" action="/analyze" enctype="multipart/form-data">
    <div class="mb-3">
      <label class="form-label">File (optional)</label>
      <input class="form-control" type="file" name="file">
    </div>
    <div class="mb-3">
      <label class="form-label">Or paste text (optional)</label>
      <textarea class="form-control" name="pasted_text" rows="8"></textarea>
    </div>
    <button class="btn btn-primary">Analyze</button>
  </form>
  <hr>
  <p class="text-muted">Note: spaCy, pdfplumber, python-docx and NLTK may be required for full functionality.</p>
</div>
</body>
</html>
"""

RESULT_HTML = """
<!doctype html>
<html>
<head><meta charset="utf-8"><title>Analysis Result</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"></head>
<body class="bg-white">
<div class="container py-4">
  <a href="/" class="btn btn-link">&larr; Back</a>
  <h1>Analysis Result</h1>
  <p><strong>Source:</strong> {{ source }}</p>

  <h4>Counts</h4>
  <ul>
    <li>Characters (with spaces): {{ stats.chars }}</li>
    <li>Characters (no spaces): {{ stats.chars_no_space }}</li>
    <li>Words: {{ stats.words }}</li>
    <li>Sentences: {{ stats.sentences }}</li>
    <li>Paragraphs: {{ stats.paragraphs }}</li>
  </ul>

  <h4>Top words (excluding stopwords)</h4>
  <ol>
  {% for w,c in top_words %}
    <li>{{ w }} — {{ c }}</li>
  {% endfor %}
  </ol>

  <h4>Word Frequency Chart</h4>
  <img src="/static/{{ chart_filename }}" class="img-fluid" alt="Bar chart">

  <h4>Named Entities</h4>
  <div class="row">
    <div class="col-md-6">
      <h5>People</h5>
      <ul>{% for e in entities.PERSON %}<li>{{ e }}</li>{% endfor %}</ul>
    </div>
    <div class="col-md-6">
      <h5>Organizations</h5>
      <ul>{% for e in entities.ORG %}<li>{{ e }}</li>{% endfor %}</ul>
    </div>
  </div>
  <div class="row">
    <div class="col-md-6">
      <h5>Places (GPE/LOC)</h5>
      <ul>{% for e in entities.GPE %}<li>{{ e }}</li>{% endfor %}</ul>
    </div>
    <div class="col-md-6">
      <h5>Dates</h5>
      <ul>{% for e in entities.DATE %}<li>{{ e }}</li>{% endfor %}</ul>
    </div>
    <div class="col-md-6">
        <h5>Summary</h5>
        <p>{{ summary }}</p>
    </div>

  </div>

  <hr>
  <p class="text-muted">Generated at {{ ts }}</p>
</div>
</body>
</html>
"""

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXT

def extract_text_from_pdf(path):
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed")
    texts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            texts.append(p.extract_text() or "")
    return "\n\n".join(texts)

def extract_text_from_docx(path):
    if docx is None:
        raise RuntimeError("python-docx is not installed")
    doc = docx.Document(path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text)

def extract_text_from_txt(path):
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def basic_stats(text):
    chars = len(text)
    chars_ns = len(text.replace(" ", "").replace("\n",""))
    # ensure NLTK punkt
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    words_list = [w for w in word_tokenize(text)]
    words = len(words_list)
    sents = sent_tokenize(text)
    sentences = len(sents)
    paragraphs = len([p for p in text.split("\n\n") if p.strip()])
    return {"chars": chars, "chars_no_space": chars_ns, "words": words, "sentences": sentences, "paragraphs": paragraphs}

def top_frequent_words(text, top_n=15):
    # ensure stopwords
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))
    # add common articles to be safe (they are in stopwords usually)
    stop_words.update(['the','a','an'])
    tokens = [t.lower() for t in word_tokenize(text) if re.match(r"^[A-Za-z0-9'-]+$", t)]
    filtered = [t for t in tokens if t not in stop_words and len(t) > 1]
    c = Counter(filtered)
    return c.most_common(top_n), c

def save_bar_chart(top_items, filename):
    # top_items: list of (word, count)
    words = [w for w,c in top_items]
    counts = [c for w,c in top_items]
    plt.figure(figsize=(8, max(3, len(words)*0.4)))
    plt.barh(words[::-1], counts[::-1])
    plt.xlabel("Frequency")
    plt.tight_layout()
    path = os.path.join(STATIC_DIR, filename)
    plt.savefig(path)
    plt.close()
    return filename

def extract_entities(text):
    # returns dict lists for PERSON, ORG, GPE, DATE
    default = {"PERSON": [], "ORG": [], "GPE": [], "DATE": []}
    if spacy is None:
        return default
    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        # try to download
        try:
            from spacy.cli import download
            download("en_core_web_sm")
            nlp = spacy.load("en_core_web_sm")
        except Exception:
            return default
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in default:
            default[ent.label_].append(ent.text)
        else:
            # map similar labels
            if ent.label_ in ("GPE","LOC"):
                default["GPE"].append(ent.text)
            elif ent.label_ == "PERSON":
                default["PERSON"].append(ent.text)
            elif ent.label_ == "ORG":
                default["ORG"].append(ent.text)
            elif ent.label_ == "DATE":
                default["DATE"].append(ent.text)
    # dedupe while preserving order
    for k in default:
        seen = set()
        out = []
        for x in default[k]:
            if x not in seen:
                seen.add(x); out.append(x)
        default[k] = out
    return default

def summarize_text(text, n=2):
    """Simple frequency-based summarization"""
    sentences = re.split(r'(?<=[.!?]) +', text)
    text_clean = re.sub(r'[^\w\s]', '', text.lower())
    words = text_clean.split()
    freq = Counter(words)
    sentence_scores = {}
    for sentence in sentences:
        sentence_clean = re.sub(r'[^\w\s]', '', sentence.lower()).split()
        score = sum(freq.get(word, 0) for word in sentence_clean)
        sentence_scores[sentence] = score
    summary = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:n]
    return " ".join(summary)

@app.route("/")
def index():
    return render_template_string(INDEX_HTML)

@app.route("/analyze", methods=["POST"])
def analyze():
    file = request.files.get("file")
    pasted = request.form.get("pasted_text", "").strip()
    text = ""

    if file and file.filename != "":
        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            return "File type not allowed. Use PDF, DOCX, or TXT."
        fp = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(fp)
        ext = filename.rsplit(".",1)[1].lower()
        try:
            if ext == "pdf":
                text = extract_text_from_pdf(fp)
            elif ext == "docx":
                text = extract_text_from_docx(fp)
            else:
                text = extract_text_from_txt(fp)
        except Exception as e:
            text = f"Error extracting file: {e}"
    if pasted:
        text = (text + "\n\n" + pasted).strip() if text else pasted

    if not text:
        return redirect(url_for("index"))
    
    summary = summarize_text(text)

    stats = basic_stats(text)
    top_n = 12
    top_list, freq = top_frequent_words(text, top_n=top_n)
    chart_fn = f"chart_{int(datetime.now().timestamp())}.png"
    save_bar_chart(top_list[:10], chart_fn)
    entities = extract_entities(text)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # present mapped entity lists with keys PERSON, ORG, GPE, DATE
    mapped = {
        "PERSON": entities.get("PERSON", []),
        "ORG": entities.get("ORG", []),
        "GPE": entities.get("GPE", []),
        "DATE": entities.get("DATE", [])
    }

    return render_template_string(RESULT_HTML,
                                  content = "Uploaded file" if file and file.filename else "Pasted text",
                                  stats = stats,
                                  top_words = top_list,
                                  chart_filename = chart_fn,
                                  entities = mapped,
                                  ts = ts,
                                  summary = summary)





@app.route("/static/<path:fn>")
def static_files(fn):
    return send_from_directory(STATIC_DIR, fn)

if __name__ == "__main__":
    app.run(debug=True, port=5000)
